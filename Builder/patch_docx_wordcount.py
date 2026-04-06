"""Insert Section 12 word-count summary before APPENDICES if not already present."""
from docx import Document

DOCX = r"c:\AI\Git\training\CIHHC-IBM-Frawmework_PPs\ci-policy-app\Builder\Governing Body Authority & Responsibilities (3).docx"

WC_LINES = [
    "12. Section Word Count Summary (Enterprise Taxonomy Reconciliation)",
    "The following word counts reflect the policy body (Sections 1–11) of GV-GB-001 as maintained in the enterprise policy system. Counts include narrative and tabular policy text for each section; appendix forms are excluded.",
    "Section\tApprox. word count\tNotes",
    "1 — Policy Header\t77\tMetadata and control fields.",
    "2 — Purpose\t94\tRegulatory basis statement.",
    "3 — Scope\t94\tApplicability (includes exclusion note).",
    "4 — Policy Statement\t369\tNumbered policy statements.",
    "5 — Definitions\t202\tDefined terms.",
    "6 — Procedures\t1,931\tAll subsections 6.1–6.5.",
    "7 — Documentation Requirements\t552\tRecords and retention.",
    "8 — Compliance & Audit\t630\tMeasurement, surveyor expectations, failure points.",
    "9 — References\t711\tFederal, CMS, OIG, and cross-walk.",
    "10 — Training & Acknowledgment\t204\tOrientation and attestation.",
    "11 — Version Control\t147\tLifecycle and revision rules.",
    "Total (Sections 1–11)\t5,010\tExcludes appendices and form templates.",
]


def main() -> None:
    doc = Document(DOCX)
    for p in doc.paragraphs:
        if "Section Word Count Summary (Enterprise Taxonomy" in p.text:
            print("Already patched; skipping.")
            return

    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "APPENDICES" or t.startswith("APPENDICES"):
            anchor_idx = i
            break
    if anchor_idx is None:
        raise SystemExit("Could not find APPENDICES anchor paragraph.")

    anchor = doc.paragraphs[anchor_idx]
    # Insert from bottom up so order is correct
    for line in reversed(WC_LINES):
        anchor.insert_paragraph_before(line)

    doc.save(DOCX)
    print("Inserted word count section before APPENDICES.")


if __name__ == "__main__":
    main()
